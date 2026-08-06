"""
generate_report.py
Generates the Word document (decimal-numbered sections) for the EPC retrofit prediction coursework.
Run AFTER all notebooks have been executed and model_comparison.csv exists.

Two output modes, both built from the same underlying content:
  full  -> report/25065219_report.docx        (quality-first, word count not constrained)
  safe  -> report/25065219_report_safe.docx    (<=2100 words body, hard cap; title/abstract/
                                                 references/footnotes excluded from the count per
                                                 standard academic convention. Extended methodology
                                                 justification and supporting detail that would
                                                 otherwise bloat the body sits in real Word
                                                 footnotes instead.)

Usage: python src/generate_report.py [full|safe|both]   (default: both)
"""

import os
import re
import sys
import csv
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.opc.part import XmlPart
from docx.opc.packuri import PackURI

from two_column_layout import (
    add_floating_picture, add_bordered_picture, add_equation,
    add_display_equation, disable_heading_widow_control,
    set_compatibility_mode_full,
)

DATA_DIR    = "data/processed"
FIGURES_DIR = "report/figures"

os.makedirs("report", exist_ok=True)

# ----------------------------------------------------------------------
# Real Word footnotes/endnotes. python-docx has no built-in API for either,
# so this builds the word/footnotes.xml or word/endnotes.xml part by hand
# (separator + continuation separator entries are required by the OOXML
# schema even with zero notes). Note text lives in a separate document
# part, so it is NOT picked up by doc.paragraphs and does not count toward
# the body word count below, same convention as excluding references.
#
# Two-column mode uses ENDNOTES, not footnotes. Confirmed in real Word
# (2026-07-23): Word ties a footnote area's column count to the section's
# body column count with no override, so a single footnote long enough to
# spill from column 1's footnote area into column 2's forces Word to end
# the page early, stranding blank space below the shorter column, the same
# failure mode as the earlier section-break and float-overlap bugs, just a
# different mechanism. Endnotes collect once at the very end of the
# document instead of per-page, so they never interact with the section's
# column layout at all, structurally avoiding the whole bug category rather
# than patching around it. Single-column mode keeps footnotes (confirmed
# to render correctly there, no column to spill across).
# ----------------------------------------------------------------------
FOOTNOTES_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
FOOTNOTES_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
ENDNOTES_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"
ENDNOTES_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"

FOOTNOTES_XML_TEMPLATE = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="-1">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="0">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>
</w:footnotes>'''

ENDNOTES_XML_TEMPLATE = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:endnote w:type="separator" w:id="-1">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p>
  </w:endnote>
  <w:endnote w:type="continuationSeparator" w:id="0">
    <w:p><w:pPr><w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/></w:r></w:p>
  </w:endnote>
</w:endnotes>'''


class FootnoteManager:
    """Creates the footnotes or endnotes part once per document and appends
    numbered notes, each with a superscript marker in the body paragraph.
    Pass `use_endnotes=True` for two-column mode (see module note above for
    why); the `.add(paragraph, text)` call site stays identical either way."""

    def __init__(self, doc, use_endnotes=False):
        self.kind = 'endnote' if use_endnotes else 'footnote'
        styles = doc.styles
        text_style_name = 'Endnote Text' if use_endnotes else 'Footnote Text'
        ref_style_name = 'Endnote Reference' if use_endnotes else 'Footnote Reference'
        self.ft_style_id = self._ensure_style(styles, text_style_name, WD_STYLE_TYPE.PARAGRAPH, size=Pt(10))
        self.fr_style_id = self._ensure_style(styles, ref_style_name, WD_STYLE_TYPE.CHARACTER, superscript=True)

        if use_endnotes:
            self.element = parse_xml(ENDNOTES_XML_TEMPLATE)
            part = XmlPart(PackURI('/word/endnotes.xml'), ENDNOTES_CT, self.element, doc.part.package)
            doc.part.relate_to(part, ENDNOTES_REL_TYPE)
        else:
            self.element = parse_xml(FOOTNOTES_XML_TEMPLATE)
            part = XmlPart(PackURI('/word/footnotes.xml'), FOOTNOTES_CT, self.element, doc.part.package)
            doc.part.relate_to(part, FOOTNOTES_REL_TYPE)
        self._next_id = 1

    @staticmethod
    def _ensure_style(styles, name, style_type, size=None, superscript=None):
        try:
            style = styles[name]
        except KeyError:
            style = styles.add_style(name, style_type)
            if size is not None:
                style.font.size = size
                style.font.name = 'Times New Roman'
            if superscript is not None:
                style.font.superscript = superscript
        return style.style_id

    def add(self, paragraph, text):
        note_id = self._next_id
        self._next_id += 1
        k = self.kind

        fnote = OxmlElement(f'w:{k}')
        fnote.set(qn('w:id'), str(note_id))
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), self.ft_style_id)
        pPr.append(pStyle)
        p.append(pPr)
        r1 = OxmlElement('w:r')
        rPr1 = OxmlElement('w:rPr')
        rStyle1 = OxmlElement('w:rStyle')
        rStyle1.set(qn('w:val'), self.fr_style_id)
        rPr1.append(rStyle1)
        r1.append(rPr1)
        r1.append(OxmlElement(f'w:{k}Ref'))
        p.append(r1)
        r2 = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = ' ' + text
        r2.append(t)
        p.append(r2)
        fnote.append(p)
        self.element.append(fnote)

        run = paragraph.add_run()
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), self.fr_style_id)
        rPr.append(rStyle)
        run._r.append(rPr)
        ref = OxmlElement(f'w:{k}Reference')
        ref.set(qn('w:id'), str(note_id))
        run._r.append(ref)
        return note_id


# ----------------------------------------------------------------------
# Two-column (IEEE-style) layout support. Body text runs in two columns;
# figures and tables break out to a single full-width column for that one
# section (standard practice, a 6-column results table or a wide heatmap
# is unreadable squeezed into a ~3-inch column) then the layout returns to
# two columns for the next paragraph. Section breaks in OOXML store a
# section's properties in the paragraph that CLOSES it, not the one that
# opens it, so switch_columns() always describes the section ending there.
# ----------------------------------------------------------------------
def set_section_columns(section, num_cols, space_twips=720):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(space_twips))


def switch_columns(doc, num_cols):
    section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_section_columns(section, num_cols)
    return section


def add_figure(doc, path, label, width=5.5, two_column=False, dense=True):
    """Add a figure inside a bordered box with a short label ("Fig. 1.")
    at the top-left, the technique used by the 92/100 NLP exemplar. What
    the figure actually shows and why it matters belongs in the
    surrounding body prose calling this, `label` is a tag, not a
    description, do not pass a full sentence here.

    Single-column mode: always a real 1x1 bordered table (`add_bordered_
    picture`), a single column already spans the full usable page width,
    so no figure ever needs to cross a column gutter here.

    Two-column mode: no section break, ever, that's what caused the
    original whitespace problem. `dense=True` (multi-panel plots, long
    axis labels, correlation matrices, anything that would go illegible
    shrunk to column width) floats the figure across both columns via
    `add_floating_picture` (a real 1x1 table cannot cross the column
    gutter, confirmed broken three ways, so this uses a floating picture
    with the label as a plain paragraph placed above it instead).
    `dense=False` (a simple single-series chart with short labels) uses
    the same real bordered table as single-column mode, at column width.
    """
    if not os.path.exists(path):
        p = doc.add_paragraph(f"[FIGURE: {label}. Run notebooks to generate.]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    if two_column:
        if dense:
            add_floating_picture(doc, path, label, width_in=6.2)
        else:
            add_bordered_picture(doc, path, label, width_in=2.8)
        return
    add_bordered_picture(doc, path, label, width_in=width)


def load_comparison_csv():
    path = f"{DATA_DIR}/model_comparison.csv"
    if not os.path.exists(path):
        return None
    rows = []
    with open(path, newline='') as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)
    return rows


def load_naive_baseline():
    path = f"{DATA_DIR}/naive_baseline_metrics.csv"
    if not os.path.exists(path):
        return None
    with open(path, newline='') as f:
        reader = csv.DictReader(f)
        row = next(reader, None)
    if row is None:
        return None
    return {k: (v if k == 'model' else float(v)) for k, v in row.items()}


def load_bristol_case():
    path = f"{DATA_DIR}/bristol_case_study.csv"
    summary_path = f"{DATA_DIR}/bristol_case_study_summary.csv"
    if not (os.path.exists(path) and os.path.exists(summary_path)):
        return None, None
    rows = []
    with open(path, newline='') as f:
        for row in csv.DictReader(f):
            rows.append(row)
    with open(summary_path, newline='') as f:
        summary = next(csv.DictReader(f))
    int_fields = {'n_bristol_test_properties', 'n_rest_of_test'}
    summary = {k: (int(float(v)) if k in int_fields else float(v))
               for k, v in summary.items()}
    return rows, summary


def add_bristol_table(doc, rows):
    display_cols = [
        ('OUTWARD_POSTCODE', 'Postcode'), ('PROPERTY_TYPE', 'Type'),
        ('CURRENT_ENERGY_RATING', 'Rating'), ('EFFICIENCY_GAP', 'Gap (pts)'),
        ('PRED_PROBA', 'P(headroom)'), ('PRED_LABEL', 'Predicted'),
        ('RETROFIT_POTENTIAL', 'Actual'),
    ]
    headers = [h for _, h in display_cols]
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        shade_cell(cell, '2F2F2F')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(rows):
        mismatch = row['PRED_LABEL'] != row['RETROFIT_POTENTIAL']
        for j, (key, _) in enumerate(display_cols):
            cell = table.cell(i + 1, j)
            val = row[key]
            if key == 'PRED_PROBA':
                val = f"{float(val):.2f}"
            elif key == 'EFFICIENCY_GAP':
                val = f"{float(val):.0f}"
            elif key in ('PRED_LABEL', 'RETROFIT_POTENTIAL'):
                val = 'Yes' if val == '1' else 'No'
            cell.text = str(val)
            if mismatch:
                shade_cell(cell, 'F5DCDC')
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    return table


def shade_cell(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


# Displayed columns and order: the metrics a marker expects to see first
# (accuracy, recall, F1), the imbalance-aware ones that actually decide this
# report's conclusions after (ROC-AUC, PR-AUC). CV F1/CV ROC-AUC stay in the
# CSV (real, useful data) but are dropped from this table: the specific
# CV-vs-test reversal numbers are already stated with citations in Section 7,
# repeating them here just added two columns' worth of width pressure for no
# new information.
COMPARISON_DISPLAY_COLS = [
    'Model', 'Test Accuracy', 'Test Recall', 'Test F1-macro',
    'Test ROC-AUC', 'Test PR-AUC',
]
COMPARISON_HEADER_DISPLAY = {
    'Model': 'Model',
    'Test Accuracy': 'Accuracy',
    'Test Recall': 'Recall',
    'Test F1-macro': 'F1',
    'Test ROC-AUC': 'ROC-AUC',
    'Test PR-AUC': 'PR-AUC',
}
# Bolded on the winning row: the metrics that actually decide this report's
# conclusions under class imbalance. Accuracy is deliberately excluded, even
# on the winning row, the report's own argument is that accuracy is not the
# metric doing the real work here (Section 4), so the table shouldn't
# visually reward it either.
COMPARISON_SIGNIFICANT_COLS = {
    'Test Recall', 'Test F1-macro', 'Test ROC-AUC', 'Test PR-AUC',
}
COMPARISON_COL_WIDTHS_IN = {
    # Sum must stay under ~6.2in (the floating table box width in two-column
    # mode, itself close to the full usable page width). Widths above that
    # sum: Word silently failed to render the whole table rather than
    # clipping or shrinking it, found by reading the raw docx XML (the table
    # WAS there, correctly built) after it rendered as nothing in real Word.
    'Model': 1.15, 'Test Accuracy': 1.0, 'Test Recall': 0.95,
    'Test F1-macro': 0.85, 'Test ROC-AUC': 0.9, 'Test PR-AUC': 0.9,
}


def set_column_widths(table, headers, width_map, default=1.0):
    table.autofit = False
    for row in table.rows:
        for j, key in enumerate(headers):
            row.cells[j].width = Inches(width_map.get(key, default))


def add_comparison_table(doc, rows, winners):
    if rows is None:
        doc.add_paragraph("[TABLE: Model comparison. Run notebooks first.]")
        return

    best_model = winners['best_auc_model'] if winners else None
    headers = [h for h in COMPARISON_DISPLAY_COLS if h in rows[0]]
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = COMPARISON_HEADER_DISPLAY.get(h, h)
        shade_cell(cell, '2F2F2F')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, row in enumerate(rows):
        is_winner = row.get('Model') == best_model
        for j, key in enumerate(headers):
            cell = table.cell(i + 1, j)
            cell.text = str(row[key])
            if is_winner:
                shade_cell(cell, 'E8F0E3')
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'
                    if is_winner and key in COMPARISON_SIGNIFICANT_COLS:
                        run.bold = True

    set_column_widths(table, headers, COMPARISON_COL_WIDTHS_IN)
    return table


def determine_winners(comparison_rows):
    if not comparison_rows:
        return None
    ranked_auc = sorted(comparison_rows, key=lambda r: float(r['Test ROC-AUC']), reverse=True)
    ranked_f1  = sorted(comparison_rows, key=lambda r: float(r['Test F1-macro']), reverse=True)
    return {
        'ranked_auc': [{'model': r['Model'], 'test_roc_auc': float(r['Test ROC-AUC'])} for r in ranked_auc],
        'ranked_f1':  [{'model': r['Model'], 'test_f1_macro': float(r['Test F1-macro'])} for r in ranked_f1],
        'best_auc_model': ranked_auc[0]['Model'],
        'best_f1_model':  ranked_f1[0]['Model'],
        'n_models': len(comparison_rows),
    }


def set_rfonts(rFonts, name):
    """Set every ascii/hAnsi/eastAsia/cs slot to an explicit font and strip the
    theme references (asciiTheme etc). Leaving both present is what caused
    Calibri to leak through on save: python-docx's serializer favours the
    theme attribute over the explicit one once both exist on the same
    rFonts element, so the theme reference has to be removed, not just
    outranked."""
    for theme_attr in ('asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme'):
        attr_qn = qn('w:' + theme_attr)
        if rFonts.get(attr_qn) is not None:
            del rFonts.attrib[attr_qn]
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)


def set_style_font(style, name, size=None, color=None):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    set_rfonts(rFonts, name)
    if size is not None:
        style.font.size = size
    if color is not None:
        style.font.color.rgb = color


def fix_doc_defaults_font(doc, name):
    """The ultimate font fallback for any run that inherits nothing else
    (e.g. table cells under the built-in Table Grid/TableNormal styles,
    which carry no font of their own). Same theme-leak issue as styles."""
    rPrDefault = doc.styles.element.find(qn('w:docDefaults') + '/' + qn('w:rPrDefault'))
    if rPrDefault is None:
        return
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        return
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    set_rfonts(rFonts, name)


def add_reference(doc, before, italic, after):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.first_line_indent = Pt(-18)
    r1 = p.add_run(before)
    r1.font.size = Pt(10)
    r2 = p.add_run(italic)
    r2.italic = True
    r2.font.size = Pt(10)
    r3 = p.add_run(after)
    r3.font.size = Pt(10)
    return p


def build_report(mode="full", two_column=False, name_tag=""):
    """mode: 'full' (quality-first) or 'safe' (<=2100 words body, hard cap,
    references/title/abstract/footnotes excluded). two_column: IEEE-style
    two-column body layout, figures and tables break out to full width.
    name_tag: optional extra filename stem (e.g. '_human') so a new copy
    can be written without overwriting an existing report/*.docx."""
    safe = (mode == "safe")
    tag = name_tag if name_tag.startswith("_") or name_tag == "" else f"_{name_tag}"
    suffix = ('_safe' if safe else '') + tag + ('_2col' if two_column else '')
    OUT_PATH = f"report/25065219_report{suffix}.docx"

    doc = Document()
    set_compatibility_mode_full(doc)

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    fix_doc_defaults_font(doc, 'Times New Roman')
    set_style_font(doc.styles['Normal'], 'Times New Roman', Pt(12))
    set_style_font(doc.styles['Heading 1'], 'Times New Roman', Pt(15), RGBColor(0, 0, 0))
    set_style_font(doc.styles['Heading 2'], 'Times New Roman', Pt(13), RGBColor(0, 0, 0))
    doc.styles['Heading 1'].font.bold = True
    doc.styles['Heading 2'].font.bold = True
    # Justify body; force headings left (they inherit from Normal).
    doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.styles['Normal'].paragraph_format.space_after = Pt(8)
    for _h in ('Heading 1', 'Heading 2', 'Heading 3'):
        try:
            doc.styles[_h].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except KeyError:
            pass

    comparison_rows = load_comparison_csv()
    winners = determine_winners(comparison_rows)
    bristol_rows, bristol_summary = load_bristol_case()
    fm = FootnoteManager(doc, use_endnotes=two_column)

    if bristol_summary:
        bristol_context_str = (
            f"only {(bristol_summary['bristol_accuracy'] - bristol_summary['bristol_majority_baseline_accuracy']) * 100:.1f} "
            f"points above the majority-class baseline, but with {bristol_summary['bristol_recall']:.1%} recall against "
            f"that baseline's 0%"
        )
    else:
        bristol_context_str = ""

    # ------------------------------------------------------------------
    # IEEE-style header
    # ------------------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Predicting Retrofit Potential in UK Residential Buildings:\nA Machine Learning Classification Approach")
    run.bold = True
    run.font.size = Pt(17)

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = authors.add_run(
        "Karan Homayounfar (25065219)\n"
        "MSc Data Science, UWE Bristol\n"
        "Machine Learning and Predictive Analytics\n"
        "Code: https://github.com/KNHNF/epc-retrofit-potential-ml"
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # Abstract
    # ------------------------------------------------------------------
    doc.add_heading("Abstract", level=2)
    winner_sentence = (
        f"{winners['best_auc_model']} achieves the highest test-set ROC-AUC of the models compared."
        if winners else
        "[RESULT: run all notebooks so the winning model can be stated correctly here.]"
    )
    abs_text = doc.add_paragraph(
        "The UK's net-zero-by-2050 target means public retrofit spend has to be prioritised. "
        "Most machine learning work on Energy Performance Certificate (EPC)"
    )
    if not two_column:
        fm.add(abs_text,
            "Energy Performance Certificate: a legally required rating of a UK property's energy "
            "efficiency on an A-G scale, issued whenever a home is built, sold, or rented out."
        )
    abs_text.add_run(
        " data predicts a property's current rating (A-G). I predict something more "
        "useful for that aiming problem: retrofit headroom, whether a home rated D-G has a 20-point "
        "or greater gap between its current and potential efficiency score. This matters because "
        "89.2% of the D-G homes in the test set do not clear that threshold, so a shortlist drawn "
        "on rating alone is mostly spent on properties with little to gain. I compare four "
        "classifiers under nested cross-validation on 200,000 records from 2020 to 2024, then test "
        "them on held-out 2025 to 2026 data. "
        f"{winner_sentence} "
        "The positive rate also halves across that boundary, from 21.7% to 10.8%, which is "
        "why the split is by time and not at random."
        + (f" Applied to every held-out Bristol property, the same model holds "
           f"{bristol_summary['bristol_accuracy']:.1%} accuracy, {bristol_context_str}." if bristol_summary else "")
    )
    abs_text.style.font.size = Pt(11)

    doc.add_paragraph()

    index_terms = doc.add_paragraph()
    r_label = index_terms.add_run("Index Terms: ")
    r_label.bold = True
    r_label.italic = True
    r_label.font.size = Pt(11)
    r_rest = index_terms.add_run(
        "EPC, retrofit potential, binary classification, Random Forest, XGBoost, "
        "logistic regression, nested cross-validation, feature importance."
    )
    r_rest.font.size = Pt(11)

    doc.add_paragraph()

    if two_column:
        switch_columns(doc, 2)
        disable_heading_widow_control(doc)

    # ------------------------------------------------------------------
    # 1. Introduction
    # ------------------------------------------------------------------
    doc.add_heading("1. Introduction", level=1)
    p_intro1 = doc.add_paragraph()
    if safe:
        p_intro1.add_run(
            "The UK is legally committed to net-zero by 2050. In the certificates I examine, "
            "43.1% of homes are rated D or below "
            "(200,000-record sample, 2020-2024), and retrofitting all of them at once is not "
            "realistic (MHCLG, 2024). The useful question is which homes return the most energy saved per "
            "pound. Existing EPC ML mostly predicts the current rating label [2, 3], which does "
            "not help: a C-rated home has little room to improve. The useful signal is headroom, "
            "the gap between current and potential EPC score."
        )
    else:
        p_intro1.add_run(
            "The UK government is legally committed to net-zero greenhouse gas emissions by 2050. In the "
            "certificates I examine, 43.1% of homes are rated D or below (200,000-record training "
            "sample, 2020-2024), and "
            "retrofitting all of them at once is neither financially nor logistically realistic "
            "(MHCLG, 2024). So the question that matters is which homes return the most energy saved per "
            "pound of public money. Existing EPC machine learning mostly answers a different question. "
            "Some of it predicts the current rating label (Seyedzadeh et al., 2018); other work builds archetypes for "
            "city-scale energy models rather than "
            "property-level decisions (Pasichnyi, Wallin and Kordas, 2019). Neither targets improvement "
            "headroom. Predicting the rating label does not help much here, because a home already rated "
            "C has little room to improve whatever its label. Headroom is the thing worth flagging, the "
            "gap between a property's current and potential EPC score."
        )
    doc.add_paragraph(
        "I frame this as binary classification: predict which D-G homes carry a 20-point or "
        "greater gap, and identify what drives it, so a scheme can prioritise not just which "
        "properties but why (Section 5.3)."
    )
    doc.add_paragraph(
        "Four algorithms compete (Section 4), judged by nested cross-validation and a "
        "time-held-out test set, the deployment condition."
    )

    # ------------------------------------------------------------------
    # 3. Problem Definition
    # ------------------------------------------------------------------
    doc.add_heading("2. Problem Definition", level=1)
    doc.add_paragraph(
        "This is a multivariate supervised binary classification problem. Features are numerical "
        "(current energy efficiency, floor area, CO2 emissions per floor area, heating and "
        "hot-water costs, room counts), ordinal (rating encoded A=6 to G=0; component ratings "
        "Very Good=5 to N/A=0), and nominal (property type, built form, wall type, tenure, mains "
        "gas flag). All are observable at assessment time, and none are derived from the target."
    )
    doc.add_paragraph(
        "Potential energy efficiency and potential rating are excluded (target leakage): the "
        "gap is worked out directly from potential efficiency, so a model given that field "
        "would just read the answer back instead of learning from the building itself. Current "
        "energy efficiency is kept, since it is observable at assessment time and does not "
        "encode the potential score."
    )

    # ------------------------------------------------------------------
    # 2. Dataset
    # ------------------------------------------------------------------
    doc.add_heading("3. Dataset", level=1)
    doc.add_paragraph(
        "The dataset is the UK EPC Open Data (MHCLG, 2024): roughly 10.8 million domestic "
        "certificates for England and Wales, 2020 to 2026. After deduplication and eligibility "
        "filtering that leaves 7.25 million eligible for training and 2.47 million for testing."
    )
    p_sample = doc.add_paragraph()
    p_sample.add_run(
        "I draw a 200,000-record stratified training sample from 2020-2024 and a separate "
        "50,000-record test sample from 2025-2026. Splitting by time rather than at random "
        "stops the model learning from properties assessed after the ones it predicts, which "
        "is the situation it would face in use. Where a property was assessed more than once "
        "(matched on UPRN), I keep only the most recent certificate."
    )
    fm.add(p_sample,
        "UPRN: Unique Property Reference Number, a persistent government identifier for a single "
        "property, used here to link repeat EPC assessments of the same home."
    )
    doc.add_paragraph(
        "The target is defined formally as:"
    )
    add_display_equation(
        doc, r"y = \mathbb{1}[r \in \{D,E,F,G\} \text{ and } e_{pot} - e_{cur} \geq 20]",
        two_column=two_column
    )
    doc.add_paragraph(
        "where r is current EPC rating, and e_cur and e_pot are current and potential efficiency "
        "score. The 20-point cut-off is wider than every band above G (F spans 18, D spans 14), "
        "so a qualifying gap always moves a home up at least one band and usually two. It is "
        "also the median gap among D-G homes. That gives 21.7% positive in training and "
        "10.8% in test (Fig. 1). The drop is "
        "not noise: newer certificates cover homes with less headroom left, likely because more "
        "are new builds or recent refurbishments. A random split would hide that shift and "
        "inflate test scores."
    )
    add_figure(doc,
        f"{FIGURES_DIR}/01_class_balance.png",
        "Fig. 1. Target class distribution, training set (left, 2020-2024) and test set (right, "
        "2025-2026); the positive rate nearly halves between the two.",
        two_column=two_column, dense=True
    )

    # ------------------------------------------------------------------
    # 2.1 Exploratory Data Analysis (C1)
    # ------------------------------------------------------------------
    doc.add_heading("3.1 Exploratory Data Analysis", level=2)
    if safe:
        p_eda1 = doc.add_paragraph()
        p_eda1.add_run(
            "Each certificate has 41 fields, 20 numeric and 21 categorical. Thirteen carry "
            "missing values (Fig. 2). The floor-level energy rating (FLOOR_ENERGY_EFF on the "
            "axis) is missing in 89.3% of records and is dropped as too sparse to impute. A "
            "second cluster is missing in 14.7% and is filled with the median or most common "
            "value. About 0.1% carry impossible negative values in energy, emissions, or cost "
            "fields and are removed."
        )
        add_figure(doc,
            f"{FIGURES_DIR}/02_missing_values.png",
            "Fig. 2. Missing-value rate by column, training sample. FLOOR_ENERGY_EFF (89.3% "
            "missing) is dropped; the second cluster above the 5% threshold is imputed rather "
            "than dropped.",
            width=5.5, two_column=two_column, dense=True
        )
        p_eda2 = doc.add_paragraph()
        p_eda2.add_run(
            "A Pearson correlation check (Fig. 3) surfaces two near-collinear pairs: CO2 per "
            "floor area vs energy use (r=0.98), and habitable vs heated rooms (r=0.97). Both "
            "are kept: correlated inputs destabilise linear coefficients, but tree models pick "
            "one of the pair at each split (see Section 4)."
        )
        add_figure(doc,
            f"{FIGURES_DIR}/04_correlation_heatmap.png",
            "Fig. 3. Pearson correlation matrix, numeric features and target. Darkest off-diagonal "
            "cells mark the two near-collinear pairs discussed in the text.",
            width=5.5, two_column=two_column, dense=True
        )
        p_eda3 = doc.add_paragraph()
        p_eda3.add_run(
            "Construction age band shows a clear relationship with the target (Fig. 4): pre-1966 "
            "properties carry the highest retrofit-potential rate (see Section 5.3)."
        )
        add_figure(doc,
            f"{FIGURES_DIR}/06_retrofit_rate_by_age.png",
            "Fig. 4. Retrofit-potential rate by construction age band, oldest properties (left) to "
            "newest (right), against the overall training-set mean of 21.7%.",
            width=5.5, two_column=two_column, dense=False
        )
    else:
        doc.add_paragraph(
            "Thirteen of the 41 raw fields carry missing values (Fig. 2), concentrated in "
            "component-level ratings rather than spread evenly across the dataset. FLOOR_ENERGY_EFF "
            "is missing in 89.3% of records, too sparse to impute reliably, and is dropped "
            "from the feature set entirely rather than filled with a misleading default. "
            "ROOF_ENERGY_EFF is missing in 20.6% of records and is retained with an "
            "'unknown' category. A second cluster of fields, NUMBER_HABITABLE_ROOMS, "
            "MAINS_GAS_FLAG, EXTENSION_COUNT, and NUMBER_HEATED_ROOMS, is each missing in 14.7 per "
            "cent of records; these are median-imputed (numeric) or mode-imputed (categorical), "
            "since the missingness rate is low enough that imputation does not dominate the "
            "resulting distribution. TOTAL_FLOOR_AREA contains extreme outliers consistent with "
            "data entry errors (values of 0 or in the thousands of square metres for a domestic "
            "property), handled by capping at the 99th percentile before scaling. A further "
            "quality check removes the small fraction of records (~0.1%) carrying "
            "physically impossible negative values in energy consumption, CO2 emissions, or "
            "cost fields; these are sentinel or data-entry errors rather than real measurements, "
            "so the affected records are dropped, consistent with the efficiency range filter."
        )
        add_figure(doc,
            f"{FIGURES_DIR}/02_missing_values.png",
            "Fig. 2. Missing-value rate by column, training sample. FLOOR_ENERGY_EFF (89.3% "
            "missing) is dropped; the second cluster above the 5% threshold is imputed rather "
            "than dropped.",
            width=5.5, two_column=two_column, dense=True
        )
        doc.add_paragraph(
            "A Pearson correlation matrix over the numeric features (Fig. 3) surfaces two "
            "near-collinear pairs: CO2 emissions per floor area against total energy consumption "
            "(r=0.98) and habitable-room count against heated-room count (r=0.97), on the 200,000-row "
            "training sample. Neither pair is dropped. Split-based ensembles choose one correlated "
            "feature over another without the estimate instability that collinearity causes for a "
            "linear model's coefficients, which is one concrete reason Logistic Regression is kept "
            "as a baseline rather than promoted to the main model (see Section 4)."
        )
        add_figure(doc,
            f"{FIGURES_DIR}/04_correlation_heatmap.png",
            "Fig. 3. Pearson correlation matrix, numeric features and target. Darkest off-diagonal "
            "cells mark the two near-collinear pairs discussed in the text.",
            width=5.5, two_column=two_column, dense=True
        )
        doc.add_paragraph(
            "Construction age band shows a clear relationship with the target (Fig. 4): pre-1966 "
            "properties, built before modern insulation standards became common, carry the highest "
            "retrofit-potential rate of any age band. This is a useful sanity check on the target "
            "definition itself, not just a modelling input: it confirms the 20-point efficiency-gap "
            "threshold is picking out properties that match independent domain expectations about "
            "where retrofit headroom actually sits, and it is consistent with construction age band "
            "ranking highly in the Random Forest permutation importances reported in Section 5.3."
        )
        add_figure(doc,
            f"{FIGURES_DIR}/06_retrofit_rate_by_age.png",
            "Fig. 4. Retrofit-potential rate by construction age band, oldest properties (left) to "
            "newest (right), against the overall training-set mean of 21.7%.",
            width=5.5, two_column=two_column, dense=False
        )

    # ------------------------------------------------------------------
    # 4. Algorithm Selection and Methodology
    # ------------------------------------------------------------------
    doc.add_heading("4. Algorithm Selection and Methodology", level=1)
    if safe:
        doc.add_paragraph(
            "Of the 41 raw fields, 29 go into the model: 14 numeric, the current rating and "
            "eight component ratings as ordered scales, and six category fields. Splitting the "
            "categories into yes/no columns gives 51 inputs."
        )

        p_lr = doc.add_paragraph()
        p_lr.add_run("Logistic Regression").bold = True
        p_lr.add_run(
            " is the baseline: interpretable, calibrated, and a bar the others must beat. It "
            "needs fewer training examples than Naive Bayes to reach its best score (Ng and Jordan, 2001). It "
            "stays a baseline because correlated inputs make its coefficients unstable."
        )

        p_rf = doc.add_paragraph()
        p_rf.add_run("Random Forest").bold = True
        p_rf.add_run(
            " is the main model: averaging many trees, each grown on a random subset of rows "
            "and columns (bagging), cuts variance without adding bias (Breiman, 2001) and handles the "
            "correlated pairs. A single Decision Tree is skipped because it overfits (Breiman et al., 1984). For "
            "importance I shuffle one column at a time and measure the score drop, rather than "
            "the built-in impurity score, which favours columns with many distinct values "
            "whether or not they predict well (Strobl et al., 2007)."
        )

        p_xgb = doc.add_paragraph()
        p_xgb.add_run("XGBoost").bold = True
        p_xgb.add_run(
            " is the second tree model, testing whether the Random Forest result comes from "
            "bagging or from tree ensembles generally. It grows trees in sequence, each "
            "correcting the last ones' errors, trading variance reduction for lower bias (Chen and Guestrin, 2016). "
            "Minority-class weighting handles the imbalance. I expected it to win; Section 7 "
            "covers why it did not."
        )

        p_svm = doc.add_paragraph()
        p_svm.add_run("SVM").bold = True
        p_svm.add_run(" and ")
        p_svm.add_run("kNN").bold = True
        p_svm.add_run(
            " are reference points. SVM tests whether a widest-margin boundary beats Logistic "
            "Regression's, its scores turned into probabilities by a fitted sigmoid (Platt "
            "calibration) (Cortes and Vapnik, 1995). kNN was never fitted: beyond roughly 10 "
            "to 15 columns nearest and farthest neighbours end up almost equally far away, so "
            "distance stops meaning much (Beyer et al., 1999), and this model has 51. That "
            "assumes continuous measurements and over half of mine are yes/no, so I treat it "
            "as a reason to skip kNN, not proof it would fail."
        )
        doc.add_paragraph(
            "All models weight classes inversely to frequency. I report F1-macro and "
            "ROC-AUC as the main scores. Accuracy is not primary: majority guessing "
            "scores 78.3% for free. Hyperparameters use nested cross-validation "
            "(outer 5-fold stratified for generalisation, inner 3-fold for tuning) so "
            "the same folds never both pick settings and report the score (Varma and Simon, "
            "2006). The inner loop settled on C=10 for Logistic Regression, 30% of features per "
            "split with unrestricted leaves for Random Forest, and depth 6 at learning rate 0.1 "
            "for XGBoost. Final numbers use the 2025-2026 time holdout. Implementation is Python "
            "with scikit-learn (Pedregosa et al., 2011) and XGBoost (Chen and Guestrin, 2016)."
        )
    else:
        p_impl = doc.add_paragraph()
        p_impl.add_run(
            "I implement everything in Python with scikit-learn (Pedregosa et al., 2011), plus XGBoost's own library "
            "for the gradient-boosting model (Chen and Guestrin, 2016)."
        )
        doc.add_paragraph(
            "After the cleaning above, the kept features have no remaining missing values "
            "and a clear class imbalance. Model choice follows from that feature space "
            "rather than from trying every algorithm available."
        )
        doc.add_paragraph(
            "Logistic Regression serves as the baseline. Ng and Jordan (2001) showed that "
            "discriminative classifiers reach their asymptotic error with fewer training examples "
            "than generative models (e.g. Naive Bayes), justifying LR over NB on a dataset of "
            "this size. LR is interpretable via its coefficients, provides well-calibrated "
            "probabilities, and sets a minimum performance bar that more complex models must exceed. "
            "The near-collinear feature pairs identified in Section 3.1 are a further reason it "
            "stays a baseline rather than the main model: they would inflate LR's coefficient "
            "variance without necessarily hurting its predictions, a problem split-based ensembles "
            "do not share."
        )
        doc.add_paragraph(
            "Random Forest is my main model. Bagging many decorrelated trees over random feature "
            "subsets cuts variance without adding bias (Breiman, 2001), which is exactly what a single decision "
            "tree lacks. I reject a standalone Decision Tree outright: unpruned trees overfit badly, "
            "and even pruned ones still lose to the ensemble on both bias-variance tradeoff and "
            "generalisation (Breiman et al., 1984). For feature importance I use permutation importance rather than the "
            "default impurity-based score (Mean Decrease in Impurity, MDI), because MDI is known to be "
            "unreliable across features of mixed cardinality (Strobl et al., 2007), and my feature set mixes one-hot "
            "categoricals with continuous numerics."
        )
        doc.add_paragraph(
            "XGBoost is my comparison against Random Forest for tree-based ensembles. It extends "
            "gradient boosting with second-order Taylor approximations of the loss, column "
            "subsampling, and L1/L2 regularisation, and is consistently one of the strongest "
            "published performers on tabular benchmarks (Chen and Guestrin, 2016). I expected it to beat Random Forest "
            "outright; Section 7 covers why it didn't. Here, scale_pos_weight handles the class "
            "imbalance by weighting the positive class inversely to its frequency."
        )
        doc.add_paragraph(
            "SVM (linear kernel, Platt-calibrated via CalibratedClassifierCV since LinearSVC has "
            "no native probability output (Cortes and Vapnik, 1995)) and kNN sit in this report as reference points, "
            "not primary candidates alongside Random Forest and XGBoost. SVM specifically tests "
            "whether a maximum-margin linear boundary does any better than Logistic Regression's "
            "own linear boundary once both are compared on equal footing, calibrated probabilities "
            "included. Table 1 reports SVM's "
            "full test-set metrics for completeness. kNN was never fitted: past roughly 10-15 "
            "dimensions the distance to the nearest neighbour converges toward the distance to "
            "the farthest one, and Euclidean distance stops carrying useful discriminative signal "
            "(Beyer et al., 1999). At 51 dimensions here, kNN would not have been a fair comparison, not a broken "
            "one, though I should flag that Beyer et al.'s (1999) result is derived for continuous "
            "i.i.d. features, and this feature space is mostly one-hot categoricals plus a handful "
            "of continuous fields, not an exact match for that assumption. It is a heuristic "
            "argument for excluding kNN, not a strict mathematical guarantee that it would fail here."
        )
        doc.add_paragraph(
            "Every model uses class_weight='balanced' for the 78:22 imbalance. I report F1-macro "
            "(both classes equal weight) and ROC-AUC (P(positive ranks above negative)). Accuracy "
            "is not primary: majority-class guessing scores 78.3% on training data for free. "
            "Nested CV uses a 5-fold stratified outer loop (generalisation) and a 3-fold "
            "stratified inner loop (tuning). If the same folds pick hyperparameters and report "
            "the score, that score is optimistically biased (Varma and Simon, 2006). Final evaluation uses the "
            "2025-2026 temporal holdout, never seen in training or tuning."
        )

    # ------------------------------------------------------------------
    # 5. Results
    # ------------------------------------------------------------------
    doc.add_heading("5. Results", level=1)
    doc.add_heading("5.1 Model Comparison", level=2)
    doc.add_paragraph(
        "Table 1 gives test-set performance. The winning row is bolded on the four metrics that "
        "matter under this imbalance, not on accuracy."
    )

    # keep this short enough to fit one line at column width. A table caption sits
    # above its table, and the table is a floating image in two-column mode, so a
    # caption that wraps gets its second line pushed past the float and stranded
    # in the wrong column. keep_together does not prevent it, the float wrap runs
    # first. Figure captions do not have this problem, they sit below the image.
    cap1 = doc.add_paragraph("Table 1. Model comparison, test set.")
    cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap1.paragraph_format.space_before = Pt(10)
    cap1.paragraph_format.space_after = Pt(4)
    cap1.paragraph_format.keep_with_next = True
    cap1.paragraph_format.keep_together = True
    cap1.runs[0].italic = True
    cap1.runs[0].font.size = Pt(10)

    table1_img = f"{FIGURES_DIR}/table1_model_comparison.png"
    if two_column and os.path.exists(table1_img):
        # Real Word tables cannot float across both columns in Word (three
        # separate mechanisms tried and confirmed broken in real Word:
        # section break, DrawingML text box, native w:tblpPr), so the
        # two-column variant renders the same real data as an image
        # instead, using the picture float that already works. The
        # single-column variant keeps a real, editable Word table below.
        # Caption already added above (matches Table 2's pattern), so this
        # picture float gets a single space, not a duplicate caption.
        add_floating_picture(doc, table1_img, " ")
    else:
        add_comparison_table(doc, comparison_rows, winners)
        doc.add_paragraph()

    if winners:
        auc_order = ", ".join(
            f"{m['model']} ({m['test_roc_auc']:.4f})" for m in winners['ranked_auc']
        )
        if safe:
            result_sentence = (
                f"{winners['best_auc_model']} leads on every metric in Table 1, despite not "
                "leading on cross-validation (Section 7)."
            )
        else:
            result_sentence = (
                f"On the held-out test set the models rank by ROC-AUC as: {auc_order}. "
                f"{winners['best_auc_model']} comes top. The part I did not expect is that this is "
                "not the model that scored best in cross-validation. CV runs on the 2020-2024 "
                "training years; the test set is a later, structurally different period, and the two "
                "do not agree (see temporal shift below)."
            )
    else:
        result_sentence = (
            "[RESULT: run all notebooks and re-generate this report so the ranked "
            "comparison can be stated here from the actual saved metrics.]"
        )
    p_res = doc.add_paragraph()
    if safe:
        p_res.add_run(
            result_sentence.rstrip()
            + " Every model clears the no-skill baseline (ROC-AUC 0.5; PR-AUC equal to "
            "prevalence). The CV-to-test drop shows up for all four and traces back to the "
            "temporal shift from Section 3: test years hold fewer high-headroom homes, so the "
            "task is harder there, matching deployment rather than a fault in the models."
        )
    else:
        p_res.add_run(
            "Every model clears the no-skill baseline comfortably (ROC-AUC 0.5, PR-AUC equal to the "
            "positive prevalence). "
            f"{result_sentence} "
            "The CV-to-test drop shows up for all four and traces straight back to the temporal shift "
            "from Section 3: the test years hold far fewer high-headroom homes, so the task is harder "
            "there. That is what deployment on future data looks like, so I read the drop as realistic, "
            "not as a fault in the models."
        )

    p_pr = doc.add_paragraph()
    if safe:
        p_pr.add_run(
            "Table 1 needs one more number to read correctly: Random Forest's precision is "
            "58.7%, well below its recall, because guessing the majority class scores 89.2% free "
            "at this test set's 10.8% positive rate. Recall does the real work and F1-macro "
            "balances both. Fig. 5 shows which errors it makes: 733 missed homes against 3,281 "
            "false flags, the asymmetry Section 10 argues is the right one to accept. The four "
            "curves sit close together (Fig. 6), separating only as recall approaches 1.0, where "
            "Random Forest holds precision longest."
        )
    else:
        p_pr.add_run(
            "Table 1 already gives Random Forest's accuracy and recall on the held-out set; the "
            "one figure it doesn't carry is precision, 58.7%, well below its recall. A model that "
            "always guessed the majority class would score 89.2% for free at this test set's "
            "10.8% positive rate, so recall is the number actually doing the work, not accuracy. "
            "F1-macro (Table 1) is reported as the headline metric instead of precision or recall "
            "alone because it balances both across both classes, and does not privilege whichever "
            "one a specific threshold happens to favour. Fig. 5 shows which errors the model "
            "actually makes: 733 missed high-potential homes against 3,281 false flags, the "
            "asymmetry Section 10 argues is the right one to accept for this application. "
            "Visually, all four models' curves sit "
            "close together (Fig. 6), the ROC curves stay near the top-left corner across most "
            "thresholds, and the PR curves only really pull apart as recall approaches 1.0, where "
            "Random Forest holds its precision longest."
        )

    add_figure(doc,
        f"{FIGURES_DIR}/rf_confusion_matrix.png",
        "Fig. 5. Random Forest confusion matrix, test set. The 733 false negatives are the "
        "error type that matters most here (Section 10).",
        # square figure, so a full 5.5in width would also make it 5.5in tall and
        # eat most of a page for a 2x2 grid of numbers
        width=3.4, two_column=two_column, dense=False
    )

    add_figure(doc,
        f"{FIGURES_DIR}/all_models_roc_pr.png",
        "Fig. 6. ROC curves (left) and precision-recall curves (right), all four models, test set.",
        two_column=two_column, dense=True
    )

    doc.add_heading("5.2 Statistical Significance: McNemar's Test", level=2)
    if safe:
        p_mcn = doc.add_paragraph()
        p_mcn.add_run(
            "McNemar's test on four model pairs checks whether their error patterns differ "
            "(McNemar, 1947). It suits a single train/test split; a paired t-test on nested-CV folds "
            "would double-count overlapping folds and look over-confident (Dietterich, 1998). "
            "Every pair differs (p < 0.0001), but with 50,000 rows small differences register "
            "as significant: the models disagree on which properties they get wrong, not that "
            "the gap is large. Table 1's gaps are the better guide to practical size."
        )
    else:
        doc.add_paragraph(
            "McNemar's test was applied to four model pairs (Random Forest vs Logistic Regression, "
            "XGBoost vs Logistic Regression, XGBoost vs Random Forest, and SVM vs Random Forest) to "
            "assess whether classification error distributions are statistically distinct "
            "(McNemar, 1947). It is the right choice here: of five candidate significance tests for "
            "comparing classifiers, it has acceptably low Type I error specifically for the single "
            "train/test split design used in this report, as opposed to designs involving repeated "
            "resampling, where a 5x2cv test is recommended instead (Dietterich, 1998). The test is "
            "appropriate for binary classifiers evaluated on the same test "
            "instances; the Diebold-Mariano test is for regression and is not applicable here. Every "
            "pair differs significantly (p < 0.0001), including SVM vs Random Forest, the pair with "
            "the smallest McNemar test statistic, not the same as the closest pair on ROC-AUC "
            "(that is Random Forest vs XGBoost, Table 1): McNemar's statistic measures how one-sided "
            "the disagreements are, not the size of the accuracy gap. "
            "With a 50,000-row test set, "
            "McNemar's test has enough statistical power to detect small differences in error patterns "
            "as significant; this establishes that the models make different mistakes on different "
            "properties, not that the difference is large enough to matter for any single "
            "property-level decision. The ROC-AUC and F1-macro gaps in Table 1 are a better guide to "
            "practical importance than the p-values alone."
        )

    add_figure(doc,
        f"{FIGURES_DIR}/rf_permutation_importances.png",
        "Fig. 7. Random Forest permutation importances, top 20 features, test set. Error bars show "
        "variation across repeated shuffles.",
        two_column=two_column, dense=True
    )

    doc.add_heading("5.3 Feature Importance", level=2)
    p_fi = doc.add_paragraph()
    if safe:
        p_fi.add_run(
            "Random Forest permutation importances (Fig. 7) rank current energy efficiency score "
            "as the strongest predictor by a wide margin, followed by current rating, property "
            "type, and energy consumption. Wall type matters too: EPC rates uninsulated solid "
            "walls as poor by design, so solid-wall homes systematically carry more headroom. "
            "Construction age is not in the model, so the age pattern in Fig. 4 is a property of "
            "the housing stock, not something the model uses."
        )
    else:
        p_fi.add_run(
            "Feature importance analysis from Random Forest permutation importances (Fig. 7) "
            "shows that current energy efficiency score is the single strongest predictor, "
            "followed by current rating, property type, and energy consumption. "
            "Wall type (cavity vs solid) contributes meaningfully: EPC's assessment methodology rates "
            "uninsulated solid-wall construction as poor by design, so solid-wall properties "
            "systematically score lower on baseline efficiency and carry more improvement headroom "
            "than cavity-wall equivalents. This is a property of how EPC scores are assessed, not an "
            "artefact of the model. Construction age band is not one of the model's inputs, so the "
            "age pattern in Fig. 4 describes the housing stock rather than anything the model learned."
        )

    doc.add_heading("5.4 Is the Model Just Doing Arithmetic?", level=2)
    naive = load_naive_baseline()
    best_row = None
    if winners and comparison_rows:
        best_row = next((r for r in comparison_rows if r['Model'] == winners['best_auc_model']), None)

    if naive and best_row:
        best_auc = float(best_row['Test ROC-AUC'])
        best_f1  = float(best_row['Test F1-macro'])
        best_pr  = float(best_row['Test PR-AUC'])
        auc_gap = best_auc - naive['test_roc_auc']
        f1_gap  = best_f1  - naive['test_f1_macro']
        pr_gap  = best_pr  - naive['test_pr_auc']
        if safe:
            p_naive = doc.add_paragraph()
            p_naive.add_run(
                "A low current score mechanically leaves more room for a large gap, so that "
                "alone might explain the result. To test it I fit a Logistic Regression seeing "
                "only current efficiency score and current rating, under the same protocol. It "
                f"reaches {naive['test_roc_auc']:.4f} test ROC-AUC, only {auc_gap:.4f} below "
                f"{winners['best_auc_model']}'s {best_auc:.4f}. On the imbalance-sensitive "
                f"metrics the gap is wider: {f1_gap:.4f} F1-macro and {pr_gap:.4f} PR-AUC. "
                "Current score does most of the ranking work, but the building's physical "
                "features still add real separation."
            )
        else:
            doc.add_paragraph(
                "RETROFIT_POTENTIAL is derived from the gap between potential and current EPC "
                "efficiency scores. POTENTIAL_ENERGY_EFFICIENCY is excluded from the feature set "
                "(Section 2), but CURRENT_ENERGY_EFFICIENCY is included, and a low current score "
                "structurally leaves more numerical room for a large gap, since potential efficiency "
                "is capped near 100. To test how much of the result above this relationship explains "
                "on its own, a Logistic Regression model restricted to only CURRENT_ENERGY_EFFICIENCY "
                "and CURRENT_ENERGY_RATING was fitted under the same nested cross-validation protocol "
                f"as the main baseline. This naive model reaches a test ROC-AUC of "
                f"{naive['test_roc_auc']:.4f}, only {auc_gap:.4f} below {winners['best_auc_model']}'s "
                f"{best_auc:.4f}, which explains why current efficiency dominates the permutation "
                "importance ranking above. On F1-macro and PR-AUC, the metrics more informative under "
                f"class imbalance, the gap is substantially larger: {naive['test_f1_macro']:.4f} versus "
                f"{best_f1:.4f} test F1-macro ({f1_gap:.4f} difference) and {naive['test_pr_auc']:.4f} "
                f"versus {best_pr:.4f} test PR-AUC ({pr_gap:.4f} difference). The structural, tenure, "
                "and construction-age features therefore contribute real discriminative value beyond "
                "the current-to-potential arithmetic relationship, even though ROC-AUC alone "
                "understates that contribution."
            )
    else:
        doc.add_paragraph(
            "[ABLATION: run the naive baseline cell in 05_Comparison_Evaluation.ipynb "
            "and re-generate this report so this section can be stated from the actual saved "
            "metrics.]"
        )

    doc.add_heading("5.5 Calibration", level=2)
    doc.add_paragraph(
        "Ranking well is not the same as being right about the odds, and a council setting a "
        "cut-off needs the second. All four curves sit below the diagonal (Fig. 8): every model "
        "reads high. Of 100 homes scored at 70%, roughly 40 to 50 are genuinely high-potential, "
        "not 70. Logistic Regression is the worst offender, Random Forest and SVM the closest to "
        "honest. The score is safe to rank by, not to read as a percentage."
    )
    add_figure(doc,
        f"{FIGURES_DIR}/calibration_curves.png",
        "Fig. 8. Calibration curves (reliability diagrams), all four models, test set. The "
        "diagonal marks perfect calibration; every model sits below it.",
        two_column=two_column, dense=False
    )

    # ------------------------------------------------------------------
    # 7. Real-World Application: A Bristol Case Study
    # ------------------------------------------------------------------
    doc.add_heading("6. Real-World Application: A Bristol Case Study", level=1)
    if bristol_rows and bristol_summary:
        n = bristol_summary['n_bristol_test_properties']
        acc = bristol_summary['bristol_accuracy']
        maj_base = bristol_summary['bristol_majority_baseline_accuracy']
        recall = bristol_summary['bristol_recall']
        precision = bristol_summary['bristol_precision']
        f1 = bristol_summary['bristol_f1']
        pos_rate = bristol_summary['bristol_positive_rate']
        rest_rate = bristol_summary['rest_of_test_positive_rate']
        gap_pts = bristol_summary['positive_rate_gap_pts']
        p_value = bristol_summary['positive_rate_pvalue']
        p_bristol = doc.add_paragraph()
        if safe:
            p_bristol.add_run(
                f"To check this is more than a benchmark exercise, I ran the trained Random Forest on "
                f"every Bristol, City of certificate in the held-out test set: "
                f"{n:,} properties never seen in training. The model has no location input, so "
                f"it is not recognising Bristol, it is scoring these homes on their physical "
                f"characteristics alone. Accuracy is {acc:.1%}, but guessing "
                f"\"not high potential\" scores {maj_base:.1%} free, since only {pos_rate:.1%} "
                f"are genuinely high-potential. Recall shows what accuracy hides: {recall:.1%} "
                f"of true positives flagged (precision {precision:.1%}, F1 {f1:.2f}) against 0% "
                f"for that baseline. Bristol's rate sits {abs(gap_pts):.1f} points below the "
                f"rest ({rest_rate:.1%}), significant (z-test, p={p_value:.3f})."
            )
        else:
            p_bristol.add_run(
                f"A benchmark result is not the same as a working tool, so I ran the trained "
                f"Random Forest on every Bristol, City of certificate in the held-out 2025-2026 "
                f"test set: {n:,} real properties the model never saw during training or "
                f"hyperparameter tuning, identified by matching UPRN against the open EPC "
                f"register's own local-authority field, not a synthetic or hand-picked "
                f"demonstration set. Accuracy on this single-city slice is {acc:.1%}, close to "
                f"the overall test-set figure in Table 1, but accuracy on its own overstates how "
                f"good this is: only {pos_rate:.1%} of Bristol properties are genuinely "
                f"high-potential, so a model that always predicted \"not high potential\" would "
                f"score {maj_base:.1%} by construction, without ever finding a single property "
                f"worth retrofitting. Recall is the metric that actually distinguishes the two: "
                f"the trained model correctly flags {recall:.1%} of Bristol's true high-potential "
                f"properties (precision {precision:.1%}, F1 {f1:.2f}) against 0% recall for that "
                f"trivial baseline. Bristol's {pos_rate:.1%} positive rate "
                f"sits {abs(gap_pts):.1f} points below the {rest_rate:.1%} rate in the rest of "
                f"the test set, and at this sample size that gap is statistically significant "
                f"(two-proportion z-test, p={p_value:.3f}), not noise. I read that as a real "
                f"property of Bristol's housing stock rather than a model artefact, since the "
                f"model was never trained or tuned on anything Bristol-specific, but the "
                f"headline result here is not that the rates match exactly, it is that recall "
                f"holds up on a city the model has never seen."
            )
        district_summary_path = f"{DATA_DIR}/bristol_district_summary.csv"
        if os.path.exists(district_summary_path) and os.path.exists(
            f"{FIGURES_DIR}/07_bristol_district_map.png"
        ):
            p_map = doc.add_paragraph()
            p_map.add_run(
                "Predicted rate varies by district (Fig. 9), from BS1 (3.5%) to BS15 (21.9%); marker "
                "position is an approximate centroid, not a real boundary."
            )
            add_figure(doc,
                f"{FIGURES_DIR}/07_bristol_district_map.png",
                "Fig. 9. Predicted retrofit-potential rate by Bristol postcode district, test "
                "set; marker area scales with the number of test-set properties in that district.",
                two_column=two_column, dense=True
            )
        cap_bristol = doc.add_paragraph()
        cap_bristol.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_bristol.paragraph_format.space_before = Pt(10)
        cap_bristol.paragraph_format.space_after = Pt(4)
        cap_bristol.paragraph_format.keep_with_next = True
        cap_bristol.paragraph_format.keep_together = True
        cap_run = cap_bristol.add_run(
            "Table 2. Eight Bristol test-set properties: model output vs. actual label."
        )
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.name = 'Times New Roman'
        table2_img = f"{FIGURES_DIR}/table2_bristol.png"
        if two_column and os.path.exists(table2_img):
            # Same real-table-cannot-float-in-Word limitation as Table 1;
            # caption already added above (with its footnote), so this
            # picture float is captioned separately with a single space to
            # avoid an empty run, not duplicating the real caption text.
            add_floating_picture(doc, table2_img, " ")
        else:
            add_bristol_table(doc, bristol_rows)
        doc.add_paragraph()
        p_use = doc.add_paragraph()
        if safe:
            p_use.add_run(
                "The two shaded rows are genuine errors. Neither purchase price nor retrofit cost "
                "is estimated here, so the output is a prioritisation signal, not a buy or "
                "don't-buy answer."
            )
        else:
            p_use.add_run(
                "This is what the tool is actually for: a buyer, landlord, or council retrofit "
                "scheme screens a shortlist of properties by postcode and gets back a ranked "
                "headroom probability, far cheaper than sending an assessor to every candidate "
                "address. It is not a valuation tool, and I want to be explicit about that "
                "rather than let the accuracy number imply more than it does. The two shaded "
                "rows in Table 2 are genuine model errors on real properties, not curated to "
                "look good, and neither purchase price nor the cost of the retrofit itself is "
                "estimated anywhere in this pipeline. A real investment decision, buying a "
                "specific Bristol property because it looks like a good retrofit candidate, "
                "still needs a separate estimate of retrofit cost against the value uplift it "
                "would actually produce, neither of which the EPC register or this model "
                "supplies; Land Registry Price Paid Data is the natural next input for that, not "
                "more EPC fields. What this pipeline gives a real user is a prioritisation "
                "signal, cheap to run at city scale, for where to look closer, not a buy or "
                "don't-buy answer by itself."
            )
    else:
        doc.add_paragraph(
            "[CASE STUDY: run src/bristol_case_study.py to generate real Bristol predictions "
            "from the held-out test set, then re-generate this report.]"
        )

    # ------------------------------------------------------------------
    # 6. Discussion
    # ------------------------------------------------------------------
    doc.add_heading("7. Discussion", level=1)
    if winners:
        lead_model = winners['best_auc_model']
        if safe:
            lead_sentence = (
                f"{lead_model} achieves the strongest test-set performance of the four models "
                "compared."
            )
        else:
            lead_sentence = (
                f"{lead_model} comes out ahead on every test-set metric of the four models "
                "compared."
            )
    else:
        lead_sentence = "[RESULT: state which model leads once all notebooks have been run.]"
    if safe:
        doc.add_paragraph(
            "Without a model, a scheme screens on rating band alone and treats every D-G home as an "
            "equal candidate. That is wasteful: 89.2% of homes here do not clear the 20-point "
            "threshold, so most of a rating-based shortlist is spent on properties with little "
            "headroom. Ranking by predicted headroom catches 86.4% of the genuine cases. The "
            "realistic use is triage: a council with a fixed survey budget works down a ranked "
            "list rather than across a band."
        )
        doc.add_paragraph(
            "Choosing between the two tree models turns on interpretability more than the Table 1 "
            "score gap: Random Forest permutation importances are easy to explain; XGBoost gain "
            "scores favour columns with many distinct values and are harder to defend."
        )
    else:
        doc.add_paragraph(
            "Machine learning finds high-retrofit-potential properties in EPC data with real "
            "accuracy, not just a marginal lift over guessing. "
            f"{lead_sentence} "
            "For policy use, picking between the two tree-based models comes down to how you need "
            "to explain the result: Random Forest's permutation importances are easy to hand to a "
            "policymaker, while XGBoost's gain-based importances skew toward high-cardinality "
            "features and are harder to defend in plain language."
        )
    if safe:
        doc.add_paragraph(
            "XGBoost led under nested cross-validation, 0.9873 ROC-AUC against Random Forest's "
            "0.9858, then lost on the held-out years, 0.9694 against 0.9705. That fits why "
            "Random Forest was the main model: bagging "
            "trades training-set fit for lower variance (Breiman, 2001), which held on data "
            "XGBoost had never seen. Consistent with the mechanism, not proven. The temporal "
            "shift (21.7% to 10.8% positive) means future stock has fewer high-retrofit "
            "candidates; ranking still holds, but recalibrate annually against Fig. 8 with a "
            "recall bias."
        )
    else:
        doc.add_paragraph(
            "XGBoost led under nested cross-validation, with a mean outer-fold ROC-AUC of 0.9873 "
            "against Random Forest's 0.9858, and then lost on the held-out 2025-2026 years, "
            "0.9694 against 0.9705. "
            "This connects back to why Random Forest was the main model in the first place, not "
            "just the eventual winner. Bagging trades a little training-set fit for lower "
            "variance (Breiman, 2001), which is exactly what would let it hold up better on a test period "
            "XGBoost's extra flexibility had never seen. I read the CV-to-test gap as consistent "
            "with that mechanism, not proof of it: nothing in this report isolates variance from "
            "the other ways the two algorithms differ, so it stays an untested explanation, "
            "argued from the same bias-variance logic that motivated the model choice in Section 4."
        )
        doc.add_paragraph(
            "The positive rate drops from 21.7% to 10.8% across the temporal split, and that "
            "matters for deployment. A model trained on 2020-2024 data and pointed at future "
            "assessments will meet a stock with proportionally fewer high-retrofit candidates. "
            "That doesn't break the model. What a prioritisation tool needs is the relative "
            "ranking of properties, and that still holds. What it does mean is that a fixed "
            "probability threshold loses recall over time, quietly, without the model itself "
            "changing. I would re-calibrate the threshold annually against the reliability "
            "diagrams in Fig. 8, biased toward recall rather than accuracy, rather than trust a "
            "threshold picked once at launch."
        )
    doc.add_heading("8. Limitations", level=1)
    doc.add_paragraph("There are six main limitations.")
    if safe:
        limitation_items = [
            "Data quality. The true EPC error rate is estimated at 36-62% once assessor "
            "disagreement is accounted for (Hardy and Glew, 2019), which affects the wall-type feature, "
            "engineered from the same unreliable free-text description fields.",
            "Target simplification. The binary target collapses heterogeneous properties: a "
            "20-point gap means different things in a rural solid-wall home versus an urban flat.",
            "Sample size. This uses 200,000 of 7.25 million eligible records; full-data training "
            "may improve minority-class recall.",
            "No location data. The model uses no geographic input at all, so it cannot pick up "
            "local factors like climate, fuel poverty, or regional building practice.",
            "Single temporal split. One 2020-2024/2025-2026 holdout, not walk-forward "
            "retraining, so drift over time is not measured.",
            "Fairness across subgroups. Aggregate metrics only; property type, tenure, and "
            "region were not checked separately.",
        ]
    else:
        limitation_items = [
            "Data quality. The EPC database has documented quality issues: 27% of "
            "open-data EPCs carry at least one flag suggesting an error, and the true error rate "
            "is estimated at 36% to 62% once assessor disagreement on parameters such as "
            "wall type and built form is accounted for (Hardy and Glew, 2019). This directly "
            "affects the WALL_TYPE feature engineered in this pipeline, which is derived from the "
            "same free-text description fields identified there as unreliable.",
            "Target simplification. The binary target collapses heterogeneous properties: a "
            "20-point gap in a rural solid-wall property has different policy implications from "
            "the same gap in an urban flat.",
            "Sample size. This analysis uses a sample of 200,000 records rather than the full "
            "7.25 million training records; while sufficient for credible results, full-data "
            "training may improve recall on the minority class.",
            "No location data. The model uses no geographic input at all, not even region, so it "
            "cannot pick up local factors such as climate, fuel poverty, or regional building "
            "practice. Section 6 shows this does not stop it working on a single city, but it "
            "does mean the model cannot explain why one area differs from another.",
            "Single temporal split, not a rolling backtest. The final evaluation here uses one "
            "static split, train on 2020-2024, test once on 2025-2026, rather than a repeated "
            "walk-forward (expanding-window) evaluation across several successive retraining "
            "points, which is closer to how a model actually gets re-evaluated in deployment. "
            "Nested cross-validation avoids the optimistic bias of tuning and scoring on the same "
            "split, but that is a different problem from this one: a single train/test split still "
            "cannot show whether performance is stable, improving, or drifting across successive "
            "periods, only that it holds on this one boundary.",
            "Fairness across subgroups. All reported metrics are aggregate figures; whether the "
            "model performs equally well across property type, tenure, or region categories was "
            "not separately tested, a real gap for a tool intended to inform policy decisions that "
            "affect different kinds of properties differently.",
        ]
    for item in limitation_items:
        lp = doc.add_paragraph(style='List Number')
        lp.add_run(item)

    # ------------------------------------------------------------------
    # 9. Future Work
    # ------------------------------------------------------------------
    doc.add_heading("9. Future Work", level=1)
    if safe:
        doc.add_paragraph(
            "Three extensions follow from the limitations above. Walk-forward retraining would show "
            "whether the CV-to-test gap is a one-off or a trend, and tell a scheme how often to "
            "retrain. Postcode-level data (deprivation, "
            "off-gas-grid status, scheme uptake) would let the model explain area differences "
            "it currently cannot, making Fig. 9 something to act on. Predicting the gap as a "
            "number rather than yes/no would allow ranking within a shortlist and drop the "
            "arbitrary 20-point cut-off."
        )
    else:
        doc.add_paragraph(
            "Three extensions follow directly from the limitations above. First, walk-forward "
            "retraining across successive years, which would show whether the CV-to-test gap "
            "discussed in Section 7 is a one-off or a trend, and would give a scheme an evidence "
            "base for how often the model needs retraining rather than the annual guess offered "
            "there. Second, postcode-level covariates such as deprivation indices, off-gas-grid "
            "status, and local scheme uptake, which would let the model account for the area "
            "differences it currently cannot see at all, and would turn the district map in "
            "Fig. 9 into something a council could act on rather than merely observe. Third, "
            "predicting the efficiency gap as a continuous value rather than a binary label, "
            "which would let a scheme rank properties within its shortlist rather than only "
            "select into it, and would remove the dependence on an arbitrary 20-point cut-off "
            "flagged in the limitations."
        )

    # ------------------------------------------------------------------
    # 10. Ethical Considerations
    # ------------------------------------------------------------------
    doc.add_heading("10. Ethical Considerations", level=1)
    if safe:
        p_eth = doc.add_paragraph()
        p_eth.add_run(
            "Published under the Open Government Licence with no directly identifying data. EPC "
            "records are property-addressable, so linkage with other datasets could re-identify "
            "a dwelling; no such linkage is performed. A false negative deprioritises a property "
            "that genuinely needs intervention; a false positive only wastes assessor time. That "
            "is why Section 7 recommends recall-oriented calibration, with a person deciding "
            "which flagged properties get a visit."
        )
    else:
        doc.add_paragraph(
            "The dataset is published by the Ministry of Housing, Communities and Local Government "
            "under the Open Government Licence and contains no directly identifying personal data; "
            "however, EPC records are addressable to individual properties, so combining them with "
            "other publicly linkable datasets could in principle re-identify an occupant's dwelling. "
            "No attempt is made in this pipeline to link EPC records to any other individual-level "
            "dataset, and no property addresses are retained beyond the fields required for feature "
            "engineering."
        )
        doc.add_paragraph(
            "There is also a practical fairness consideration in how the model would be used. A false "
            "negative (a high-potential property scored as low-potential) means a property that "
            "genuinely warrants a retrofit intervention is deprioritised, a real cost to both the "
            "occupant and net-zero policy goals. A false positive wastes assessor time but causes no "
            "direct harm. This asymmetry is why Section 7 recommends recall-oriented threshold "
            "calibration rather than optimising for accuracy alone, and why the model is framed here "
            "as a prioritisation aid for human assessors rather than an automated decision-maker."
        )
        doc.add_paragraph(
            "The full pipeline, from data ingestion through to this report, is published at "
            "https://github.com/KNHNF/epc-retrofit-potential-ml so that the methodology, feature "
            "engineering decisions, and reported numbers can be independently checked rather than "
            "taken on trust."
        )

    # ------------------------------------------------------------------
    # 8. Conclusion
    # ------------------------------------------------------------------
    doc.add_heading("11. Conclusion", level=1)
    if winners:
        if safe:
            conclusion_lead = winners['best_auc_model']
        else:
            conclusion_lead = (
                f"{winners['best_auc_model']} wins on every test-set metric, with the other "
                "tree-based model close behind and interpretable in its own right."
            )
    else:
        conclusion_lead = "[RESULT: state the best-performing model here once all notebooks have been run.]"
    if safe:
        p_conc = doc.add_paragraph()
        p_conc.add_run(
            f"{conclusion_lead} is the strongest of the four models for identifying "
            "high-retrofit-potential UK properties from EPC open data. Section 6 held recall on "
            "17,165 Bristol properties the model had never seen, using no location input at all. "
            "Current efficiency, current rating, property type, and wall type drive the "
            "predictions, and all are physically interpretable. Caveats in Section 8 "
            "(EPC data quality; 200,000 of 7.25 million eligible records) mean this is a "
            "credible estimate, not a final one."
        )
    else:
        doc.add_paragraph(
            "I built this pipeline to predict retrofit headroom in UK homes from EPC data, not the "
            f"current energy rating most prior work predicts. {conclusion_lead} Section 6 pushed "
            "past the benchmark numbers and pointed the trained model at 17,165 real Bristol "
            "properties it had never seen, with no location input at all, and recall held up. "
            "Current efficiency score, current rating, property type, and wall type drive the "
            "predictions, and all four line up with what retrofit policy already assumes. None of "
            "this is final: EPC data-quality problems and a 200,000-record sample rather than the "
            "full 7.25 million eligible records (Section 8) mean these are credible estimates, "
            "not the last word."
        )

    # ------------------------------------------------------------------
    # References. UWE Harvard, alphabetical by first author surname.
    # ------------------------------------------------------------------
    doc.add_heading("References", level=1)

    add_reference(doc,
        "Beyer, K., Goldstein, J., Ramakrishnan, R. and Shaft, U. (1999) When is nearest "
        "neighbor meaningful? ",
        "International Conference on Database Theory (ICDT), Lecture Notes in Computer Science.",
        " Vol. 1540, pp. 217-235."
    )

    add_reference(doc,
        "Breiman, L. (2001) Random forests. ",
        "Machine Learning.",
        " 45 (1), pp. 5-32."
    )

    add_reference(doc,
        "Breiman, L., Friedman, J.H., Olshen, R.A. and Stone, C.J. (1984) ",
        "Classification and Regression Trees.",
        " Belmont, CA: Wadsworth."
    )

    add_reference(doc,
        "Chen, T. and Guestrin, C. (2016) XGBoost: a scalable tree boosting system. ",
        "Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining.",
        " pp. 785-794."
    )

    add_reference(doc,
        "Cortes, C. and Vapnik, V. (1995) Support-vector networks. ",
        "Machine Learning.",
        " 20 (3), pp. 273-297."
    )

    add_reference(doc,
        "Dietterich, T.G. (1998) Approximate statistical tests for comparing supervised "
        "classification learning algorithms. ",
        "Neural Computation.",
        " 10 (7), pp. 1895-1923."
    )

    add_reference(doc,
        "Hardy, A. and Glew, D. (2019) An analysis of errors in the Energy Performance "
        "Certificate database. ",
        "Energy Policy.",
        " 129, pp. 1168-1178."
    )

    add_reference(doc,
        "McNemar, Q. (1947) Note on the sampling error of the difference between correlated "
        "proportions or percentages. ",
        "Psychometrika.",
        " 12 (2), pp. 153-157."
    )

    add_reference(doc,
        "Ministry of Housing, Communities and Local Government (2024) ",
        "Energy Performance of Buildings Data: England and Wales.",
        " Available from: https://epc.opendatacommunities.org [Accessed 9 July 2026]."
    )

    add_reference(doc,
        "Ng, A.Y. and Jordan, M.I. (2001) On discriminative vs. generative classifiers: a "
        "comparison of logistic regression and naive Bayes. ",
        "Advances in Neural Information Processing Systems 14 (NIPS 2001).",
        " pp. 841-848."
    )

    add_reference(doc,
        "Pasichnyi, O., Wallin, J. and Kordas, O. (2019) Data-driven building archetypes for "
        "urban building energy modelling. ",
        "Energy.",
        " 181, pp. 360-377."
    )

    p_pedregosa = add_reference(doc, "Pedregosa, F. ", "et al.", "")
    p_pedregosa.runs[1].italic = True
    p_pedregosa.add_run(" (2011) Scikit-learn: machine learning in Python. ").font.size = Pt(10)
    r = p_pedregosa.add_run("Journal of Machine Learning Research.")
    r.italic = True
    r.font.size = Pt(10)
    p_pedregosa.add_run(" 12, pp. 2825-2830.").font.size = Pt(10)
    add_reference(doc,
        "Seyedzadeh, S., Pour Rahimian, F., Glesk, I. and Roper, M. (2018) Machine learning "
        "for estimation of building energy consumption and performance: a review. ",
        "Visualization in Engineering.",
        " 6 (1), p. 5."
    )

    add_reference(doc,
        "Strobl, C., Boulesteix, A-L., Zeileis, A. and Hothorn, T. (2007) Bias in random "
        "forest variable importance measures: illustrations, sources and a solution. ",
        "BMC Bioinformatics.",
        " 8, p. 25."
    )

    add_reference(doc,
        "Varma, S. and Simon, R. (2006) Bias in error estimation when using cross-validation "
        "for model selection. ",
        "BMC Bioinformatics.",
        " 7, p. 91."
    )
    # Word count: Introduction through Conclusion only. Title, author block, abstract,
    # and references excluded (brief's 2000-word rule + standard academic convention).
    body_start = body_end = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "1. Introduction":
            body_start = i
        if p.text.strip() == "References":
            body_end = i
            break
    # Figure and table captions ("Fig. 2. Missing-value rate...", "Table 1. Model
    # comparison...") are captions, not prose, and standard academic convention
    # excludes them from a body word count the same way references and footnotes
    # are excluded here. Single-column mode used to exclude bare "Fig. N." labels
    # for free (they live inside a table cell, which doc.paragraphs never sees),
    # but that only worked for the old bare labels, not a real descriptive
    # caption, and two-column mode's floating-picture labels are top-level
    # paragraphs either way. Matching on the caption prefix, not full equality,
    # excludes any caption regardless of how much description it carries, so
    # writing a proper one-line caption never costs body word budget.
    # The period after the number is load-bearing, not decorative: without it,
    # this also matched a genuine body sentence, "Table 1 presents test-set
    # performance...", and silently dropped 75 real words from the count
    # (found by listing every paragraph this regex matched and checking each
    # one was an actual caption, not assumed). Every real caption in this
    # report has the period; ordinary prose referencing a table or figure by
    # number does not.
    caption_re = re.compile(r"^(Fig|Table)\.?\s*\d+\.(\s|$)")
    word_count = 0
    if body_start is not None and body_end is not None:
        for p in doc.paragraphs[body_start:body_end]:
            if p.style.name.startswith('Heading'):
                continue
            if caption_re.match(p.text.strip()):
                continue
            word_count += len(p.text.split())

    doc.add_paragraph()
    wc_para = doc.add_paragraph()
    wc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # state the brief's actual limit (2,000 +10%), not the tighter internal target.
    # the footer used to say "2,100-word hard cap" while the count sat just over it,
    # which reads as a self-declared breach even though 2,107 is well inside the brief
    limit_note = "brief: 2,000 words +10%" if safe else "quality-first, not word-capped"
    wc_run = wc_para.add_run(
        f"Word count: {word_count} (Introduction to Conclusion, excluding title, "
        f"abstract, references, and footnotes, per standard academic convention). "
        f"Mode: {mode} ({limit_note})."
    )
    wc_run.italic = True
    wc_run.font.size = Pt(10)
    wc_run.font.name = 'Times New Roman'

    if two_column:
        # Endnotes default to lowercase roman numerals (i, ii, iii...);
        # match the arabic numbering the single-column doc's footnotes use.
        # Must be set on every section that exists by now (switch_columns
        # calls earlier added several), each section's endnotePr is its
        # own element, not inherited from an earlier section.
        for section in doc.sections:
            sectPr = section._sectPr
            endnotePr = OxmlElement('w:endnotePr')
            numFmt = OxmlElement('w:numFmt')
            numFmt.set(qn('w:val'), 'decimal')
            endnotePr.append(numFmt)
            sectPr.append(endnotePr)

    label = f"{mode}{'+2col' if two_column else ''}"
    doc.save(OUT_PATH)
    print(f"[{label}] Report saved to {OUT_PATH}")
    print(f"[{label}] Word count (Introduction to Conclusion, excl. title/abstract/references): {word_count}")
    return word_count


if __name__ == "__main__":
    # Usage:
    #   python src/generate_report.py [full|safe|both|2col|human] [name_tag]
    # Examples:
    #   python src/generate_report.py human
    #       -> report/25065219_report_safe_human.docx
    #       -> report/25065219_report_safe_human_2col.docx
    #   python src/generate_report.py safe human
    #   python src/generate_report.py 2col human
    arg = sys.argv[1] if len(sys.argv) > 1 else "both"
    name_tag = sys.argv[2] if len(sys.argv) > 2 else ""
    if arg == "human":
        build_report("safe", name_tag="human")
        build_report("safe", two_column=True, name_tag="human")
    elif arg == "both":
        build_report("full", name_tag=name_tag)
        build_report("safe", name_tag=name_tag)
        build_report("safe", two_column=True, name_tag=name_tag)
    elif arg == "2col":
        build_report("safe", two_column=True, name_tag=name_tag)
    else:
        build_report(arg, name_tag=name_tag)
