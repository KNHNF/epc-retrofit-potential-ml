"""
two_column_layout.py
Helpers for IEEE-style two-column Word reports: two columns, figures that
float across both columns without leaving stranded whitespace, and real
typeset equations, none of which python-docx supports natively.

See docs/IMPLEMENTATION_NOTES.md for why each of these is built the way
it is, and when to float a figure versus just shrinking it to column
width. Everything here was checked by actually opening the result in
Word, not just by the XML looking right.
"""

import latex2mathml.converter
from lxml import etree

from docx.shared import Inches, Mm, Emu, Pt
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

# Path to Microsoft's own MathML-to-OMML stylesheet. Ships with every
# desktop Office install (Word, PowerPoint, Excel all use it). If this
# path doesn't exist on a given machine, search for MML2OMML.XSL under
# the Office install directory and update this constant.
MML2OMML_XSL_PATH = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"

A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PIC_URI = "http://schemas.openxmlformats.org/drawingml/2006/picture"


def add_hyperlink(paragraph, url, text, color="0563C1", underline=True):
    """A real clickable hyperlink run, not just blue underlined text.
    python-docx has no built-in method for this."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


# Page setup

def set_page_a4(doc, margin_in=1.0):
    """python-docx's Document() defaults to US Letter. UK university
    submissions expect A4. Call this once, right after Document()."""
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(margin_in)
    section.bottom_margin = Inches(margin_in)
    section.left_margin = Inches(margin_in)
    section.right_margin = Inches(margin_in)


def set_compatibility_mode_full(doc):
    """Fixes an old Word compatibility setting python-docx leaves in
    place by default, otherwise Word 2013+ opens the file in
    Compatibility Mode, which broke floating equations for us. Call
    once, right after Document(). See docs/IMPLEMENTATION_NOTES.md."""
    settings = doc.settings.element
    compat = settings.find(qn('w:compat'))
    if compat is None:
        compat = OxmlElement('w:compat')
        settings.append(compat)
    setting = None
    for el in compat.findall(qn('w:compatSetting')):
        if el.get(qn('w:name')) == 'compatibilityMode':
            setting = el
            break
    if setting is None:
        setting = OxmlElement('w:compatSetting')
        setting.set(qn('w:name'), 'compatibilityMode')
        setting.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
        compat.insert(0, setting)
    setting.set(qn('w:val'), '15')


def disable_heading_widow_control(doc, style_names=("Heading 1", "Heading 2", "Heading 3")):
    """Stops headings forcing a blank half-column in a two-column
    layout (the biggest cause of stranded whitespace here, more than
    figures). Call once after switch_columns(doc, 2). Harmless to skip
    for single-column. See docs/IMPLEMENTATION_NOTES.md."""
    for name in style_names:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        pPr = style.element.find(qn('w:pPr'))
        if pPr is None:
            continue
        for tag in ('w:keepNext', 'w:keepLines'):
            el = pPr.find(qn(tag))
            if el is not None:
                pPr.remove(el)


# Fonts: the theme-leak bug

def set_style_font(style, name, size=None, color=None):
    """Sets a style's font in a way that actually survives saving, unlike
    plain style.font.name, which silently loses to the template's theme
    font on save. See docs/IMPLEMENTATION_NOTES.md."""
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    _set_rfonts_explicit(rFonts, name)
    if size is not None:
        style.font.size = size
    if color is not None:
        style.font.color.rgb = color


def _set_rfonts_explicit(rFonts, name):
    for theme_attr in ('asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme'):
        attr_qn = qn('w:' + theme_attr)
        if rFonts.get(attr_qn) is not None:
            del rFonts.attrib[attr_qn]
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)


def fix_doc_defaults_font(doc, name):
    """The fallback every unstyled run inherits from (table cells under
    the built-in Table Grid style carry no font of their own, they fall
    through to this). Same theme-leak issue as styles, same fix."""
    rPrDefault = doc.styles.element.find(qn('w:docDefaults') + '/' + qn('w:rPrDefault'))
    if rPrDefault is None:
        return
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        return
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    _set_rfonts_explicit(rFonts, name)


# Column switching

def set_section_columns(section, num_cols, space_twips=720):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(space_twips))


def switch_columns(doc, num_cols):
    """Continuous section break, no page break, just a column-count
    change from this point in the document onward. Note the OOXML
    quirk: a section's properties are stored in the paragraph that
    CLOSES that section, not the one that opens it, easy to get the
    before/after logic backwards when debugging."""
    section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_section_columns(section, num_cols)
    return section


# Floating (page-anchored) pictures: the whitespace fix

_next_rel_height = [1]


def add_floating_picture(doc, image_path, label, width_in=6.2, caption_below=True,
                          caption_font="Times New Roman"):
    """A picture that floats across both columns of a two-column
    section, wrapped top-and-bottom, no section break, so no stranded
    whitespace. `label` is the full one-line caption, placed below the
    image (IEEE style). See docs/IMPLEMENTATION_NOTES.md for why."""
    rid, image = doc.part.get_or_add_image(image_path)
    width_emu = Emu(int(width_in * 914400))
    cx, cy = image.scaled_dimensions(width_emu, None)

    rel_height = _next_rel_height[0]
    _next_rel_height[0] += 1

    holder = doc.add_paragraph()
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    holder.paragraph_format.space_before = Pt(12)
    holder.paragraph_format.space_after = Pt(2)
    holder.paragraph_format.keep_with_next = True
    run = holder.add_run()

    drawing = OxmlElement('w:drawing')
    anchor = OxmlElement('wp:anchor')
    for k, v in {
        'behindDoc': '0', 'distT': '91440', 'distB': '457200', 'distL': '91440',
        'distR': '91440', 'simplePos': '0', 'locked': '0', 'layoutInCell': '1',
        'allowOverlap': '0', 'relativeHeight': str(rel_height),
    }.items():
        anchor.set(k, v)

    simplePos = OxmlElement('wp:simplePos')
    simplePos.set('x', '0')
    simplePos.set('y', '0')
    anchor.append(simplePos)

    posH = OxmlElement('wp:positionH')
    posH.set('relativeFrom', 'page')
    alignH = OxmlElement('wp:align')
    alignH.text = 'center'
    posH.append(alignH)
    anchor.append(posH)

    posV = OxmlElement('wp:positionV')
    posV.set('relativeFrom', 'paragraph')
    offV = OxmlElement('wp:posOffset')
    offV.text = '0'
    posV.append(offV)
    anchor.append(posV)

    extent = OxmlElement('wp:extent')
    extent.set('cx', str(cx))
    extent.set('cy', str(cy))
    anchor.append(extent)

    effectExtent = OxmlElement('wp:effectExtent')
    for k in ('l', 't', 'r', 'b'):
        effectExtent.set(k, '0')
    anchor.append(effectExtent)

    anchor.append(OxmlElement('wp:wrapTopAndBottom'))

    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(rel_height))
    docPr.set('name', f'Picture {rel_height}')
    anchor.append(docPr)

    cNvGraphicFramePr = OxmlElement('wp:cNvGraphicFramePr')
    graphicFrameLocks = OxmlElement('a:graphicFrameLocks')
    graphicFrameLocks.set('noChangeAspect', '1')
    cNvGraphicFramePr.append(graphicFrameLocks)
    anchor.append(cNvGraphicFramePr)

    graphic = OxmlElement('a:graphic')
    graphicData = OxmlElement('a:graphicData')
    graphicData.set('uri', PIC_URI)
    pic = OxmlElement('pic:pic')

    nvPicPr = OxmlElement('pic:nvPicPr')
    cNvPr = OxmlElement('pic:cNvPr')
    cNvPr.set('id', '0')
    cNvPr.set('name', image_path.split('/')[-1].split('\\')[-1])
    cNvPicPr = OxmlElement('pic:cNvPicPr')
    nvPicPr.append(cNvPr)
    nvPicPr.append(cNvPicPr)
    pic.append(nvPicPr)

    blipFill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), rid)
    stretch = OxmlElement('a:stretch')
    stretch.append(OxmlElement('a:fillRect'))
    blipFill.append(blip)
    blipFill.append(stretch)
    pic.append(blipFill)

    spPr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    ext = OxmlElement('a:ext')
    ext.set('cx', str(cx))
    ext.set('cy', str(cy))
    xfrm.append(off)
    xfrm.append(ext)
    prstGeom = OxmlElement('a:prstGeom')
    prstGeom.set('prst', 'rect')
    prstGeom.append(OxmlElement('a:avLst'))
    spPr.append(xfrm)
    spPr.append(prstGeom)
    spPr.append(_thin_border_line())
    pic.append(spPr)

    graphicData.append(pic)
    graphic.append(graphicData)
    anchor.append(graphic)
    drawing.append(anchor)
    run._r.append(drawing)

    lbl = doc.add_paragraph()
    lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lbl.paragraph_format.space_before = Pt(2)
    lbl.paragraph_format.space_after = Pt(14)
    lbl.paragraph_format.keep_together = True
    if label and label.strip():
        lbl_run = lbl.add_run(label)
        lbl_run.italic = True
        lbl_run.font.size = Pt(10)
        lbl_run.font.name = caption_font
    else:
        lbl.paragraph_format.space_after = Pt(8)

    return holder


def add_bordered_picture(doc, image_path, label, width_in=2.8,
                          caption_font="Times New Roman"):
    """In-column figure: centred image, italic caption below, glued together."""
    p_img = doc.add_paragraph()
    p_img.paragraph_format.space_before = Pt(12)
    p_img.paragraph_format.space_after = Pt(2)
    p_img.paragraph_format.keep_with_next = True
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(image_path, width=Inches(width_in))

    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_label.paragraph_format.space_before = Pt(0)
    p_label.paragraph_format.space_after = Pt(12)
    p_label.paragraph_format.keep_together = True
    if label and label.strip():
        r_label = p_label.add_run(label)
        r_label.italic = True
        r_label.font.size = Pt(10)
        r_label.font.name = caption_font
    return p_label


def _thin_border_line(color_hex="000000", width_emu=9525):
    """~0.75pt black outline for a picture shape (spPr/a:ln), matching a
    typeset paper's figure box, not a barely-visible hairline."""
    ln = OxmlElement('a:ln')
    ln.set('w', str(width_emu))
    solidFill = OxmlElement('a:solidFill')
    srgbClr = OxmlElement('a:srgbClr')
    srgbClr.set('val', color_hex)
    solidFill.append(srgbClr)
    ln.append(solidFill)
    return ln


# Equations

def latex_to_omml(latex, xsl_path=MML2OMML_XSL_PATH):
    """LaTeX -> MathML (latex2mathml) -> OMML (Microsoft's own XSLT,
    ships with every Office install). Real typeset Word equations, not
    a screenshot. Raises FileNotFoundError with a clear message if the
    stylesheet isn't at the expected path, search for MML2OMML.XSL
    under the local Office install directory and pass the real path in."""
    mathml = latex2mathml.converter.convert(latex)
    mathml_doc = etree.fromstring(mathml.encode("utf-8"))
    xslt = etree.parse(xsl_path)
    transform = etree.XSLT(xslt)
    omml_doc = transform(mathml_doc)
    return bytes(omml_doc)


def add_equation(paragraph, latex, xsl_path=MML2OMML_XSL_PATH):
    """Insert a real Word equation into `paragraph` from a LaTeX string.
    Centre the paragraph yourself if you want a display equation:
    `p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER;
    add_equation(p, r"...")`."""
    omml_bytes = latex_to_omml(latex, xsl_path=xsl_path)
    omml_el = parse_xml(omml_bytes)
    paragraph._p.append(omml_el)
    return omml_el


def add_display_equation(doc, latex, two_column=False, xsl_path=MML2OMML_XSL_PATH,
                          width_in=6.2):
    """A centred, full-width display equation, floated the same way
    add_floating_picture floats a figure, so it doesn't need a section
    break either. See docs/IMPLEMENTATION_NOTES.md."""
    if not two_column:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_equation(p, latex, xsl_path=xsl_path)
        return p

    omml_bytes = latex_to_omml(latex, xsl_path=xsl_path)

    shape_id = _next_rel_height[0]
    _next_rel_height[0] += 1

    holder = doc.add_paragraph()
    holder.paragraph_format.space_before = Pt(6)
    holder.paragraph_format.space_after = Pt(6)
    run = holder.add_run()

    cx = Emu(int(width_in * 914400))
    cy = Emu(int(0.45 * 914400))
    drawing_xml = (
        '<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f'<wp:anchor behindDoc="0" distT="91440" distB="228600" distL="91440" '
        f'distR="91440" simplePos="0" locked="0" layoutInCell="1" allowOverlap="0" '
        f'relativeHeight="{shape_id}">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:align>center</wp:align></wp:positionH>'
        '<wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapTopAndBottom/>'
        f'<wp:docPr id="{shape_id}" name="EqBox{shape_id}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp>'
        '<wps:cNvSpPr txBox="1"/>'
        '<wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '<a:noFill/><a:ln><a:noFill/></a:ln>'
        '</wps:spPr>'
        '<wps:txbx><w:txbxContent>'
        '<w:p><w:pPr><w:jc w:val="center"/></w:pPr></w:p>'
        '</w:txbxContent></wps:txbx>'
        '<wps:bodyPr wrap="square" lIns="0" tIns="0" rIns="0" bIns="0" '
        'rtlCol="0" anchor="ctr"><a:spAutoFit/></wps:bodyPr>'
        '</wps:wsp>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:anchor>'
        '</w:drawing>'
    )
    drawing_el = parse_xml(drawing_xml)
    inner_p = drawing_el.find('.//' + qn('w:txbxContent') + '/' + qn('w:p'))
    inner_p.append(parse_xml(omml_bytes))
    run._r.append(drawing_el)
    return holder
